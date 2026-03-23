import io
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.models.report_schema import PrepLensReportV2

ASSETS = Path(__file__).resolve().parent.parent / "assets"
LOGO_PATH = ASSETS / "logo.png"

INDIGO = RGBColor(79, 70, 229)
DARK = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
EMERALD = RGBColor(5, 150, 105)
AMBER = RGBColor(217, 119, 6)
RED = RGBColor(220, 38, 38)
LIGHT_BORDER = RGBColor(203, 213, 225)


def _score_color(score: float, max_val: float = 10) -> RGBColor:
    ratio = score / max_val
    if ratio >= 0.7:
        return EMERALD
    elif ratio >= 0.4:
        return AMBER
    return RED


def _severity_color(severity: str) -> RGBColor:
    return {"high": RED, "medium": AMBER, "low": EMERALD}.get(severity, MUTED)


def generate_docx(report: PrepLensReportV2) -> bytes:
    """Generate a professionally styled DOCX from a PrepLensReportV2."""
    doc = Document()

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

    meta = report.input_summary

    # ── HEADER WITH LOGO ──
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.8))
        _add_spacing(p, before=0, after=4)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Interview Strategy Dossier")
    run.font.size = Pt(22)
    run.font.color.rgb = INDIGO
    run.bold = True
    _add_spacing(title, before=0, after=2)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{meta.company_name}  ·  {meta.job_title}")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    run.bold = True
    _add_spacing(sub, before=0, after=2)

    # Meta line
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"Generated {report.generated_at[:10]}  ·  Report ID: {report.report_id[:8]}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    _add_spacing(meta_p, before=0, after=8)

    # Divider
    _add_divider(doc)

    # ── SCORE SUMMARY TABLE ──
    pr = report.pursuit_recommendation
    rec_label = pr.recommendation.replace("_", " ").title()

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, (label, value, color) in enumerate([
        ("Overall Fit", f"{pr.overall_fit_score}/10", _score_color(pr.overall_fit_score)),
        ("Confidence", f"{pr.confidence_score}/10", _score_color(pr.confidence_score)),
        ("Recommendation", rec_label, INDIGO),
    ]):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.font.color.rgb = color
        run.bold = True

        cell2 = table.cell(1, i)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = MUTED

    # Light background for score table
    for row in table.rows:
        for cell in row.cells:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): "F1F5F9",
                qn("w:val"): "clear",
            })
            shading.append(shading_elm)

    doc.add_paragraph()

    # ── SECTIONS ──

    # 1. Executive Summary
    _section(doc, "1  Executive Summary")
    _bold_para(doc, report.executive_summary.headline)
    doc.add_paragraph(report.executive_summary.summary)
    for t in report.executive_summary.top_takeaways:
        doc.add_paragraph(t, style="List Bullet")

    # 2. Pursuit Recommendation
    _section(doc, "2  Pursuit Recommendation")
    _kv(doc, "Ideal Candidate", pr.ideal_candidate_profile)
    _kv(doc, "Your Outlook", pr.candidate_outlook)
    _sub(doc, "Reasoning")
    for reason in pr.reasoning:
        doc.add_paragraph(reason, style="List Bullet")

    # 3. Company Snapshot
    _section(doc, "3  Company Snapshot")
    cs = report.company_snapshot
    doc.add_paragraph(cs.company_overview)
    if cs.business_model:
        _kv(doc, "Business Model", cs.business_model)
    if cs.product_focus:
        _kv(doc, "Product Focus", cs.product_focus)
    if cs.company_stage != "unknown":
        _kv(doc, "Stage / Risk", f"{cs.company_stage.title()} / {cs.risk_level.title()}")
    if cs.engineering_signals:
        _sub(doc, "Engineering Signals")
        for s in cs.engineering_signals:
            doc.add_paragraph(s, style="List Bullet")

    # 4. Role Snapshot
    _section(doc, "4  Role Snapshot")
    rs = report.role_snapshot
    doc.add_paragraph(rs.role_summary)
    if rs.seniority_signal != "unknown":
        _kv(doc, "Seniority", rs.seniority_signal.title())
    if rs.primary_stack_signals:
        _kv(doc, "Stack", ", ".join(rs.primary_stack_signals))
    if rs.success_profile:
        _kv(doc, "Success Profile", rs.success_profile)

    # 5. Hiring Priorities
    _section(doc, "5  Hiring Priorities")
    for hp in report.hiring_priorities:
        p = doc.add_paragraph()
        score_run = p.add_run(f"[{hp.importance_score}/5] ")
        score_run.bold = True
        score_run.font.color.rgb = _score_color(hp.importance_score, 5)
        p.add_run(hp.priority).bold = True
        _detail(doc, hp.why_it_matters)

    # 6. Fit Analysis
    _section(doc, "6  Fit Analysis")
    if report.fit_analysis.strong_fit_signals:
        _sub(doc, "Strong Fit")
        for s in report.fit_analysis.strong_fit_signals:
            p = doc.add_paragraph()
            score_run = p.add_run(f"[{s.strength_score}/10] ")
            score_run.bold = True
            score_run.font.color.rgb = _score_color(s.strength_score)
            p.add_run(s.signal).bold = True
            _detail(doc, s.evidence)
    if report.fit_analysis.partial_fit_signals:
        _sub(doc, "Partial Fit")
        for s in report.fit_analysis.partial_fit_signals:
            p = doc.add_paragraph()
            score_run = p.add_run(f"[{s.strength_score}/10] ")
            score_run.bold = True
            score_run.font.color.rgb = _score_color(s.strength_score)
            p.add_run(s.signal).bold = True
            _detail(doc, s.evidence)
    if report.fit_analysis.missing_or_weak_signals:
        _sub(doc, "Gaps")
        for s in report.fit_analysis.missing_or_weak_signals:
            p = doc.add_paragraph()
            sev_run = p.add_run(f"[{s.gap_severity.upper()}] ")
            sev_run.bold = True
            sev_run.font.color.rgb = _severity_color(s.gap_severity)
            p.add_run(s.signal).bold = True
            _detail(doc, s.why_it_matters)

    # 7. Concerns & Mitigation
    _section(doc, "7  Concerns & Mitigation")
    for c in report.concerns_and_mitigation:
        p = doc.add_paragraph()
        sev_run = p.add_run(f"[{c.severity.upper()}] ")
        sev_run.bold = True
        sev_run.font.color.rgb = _severity_color(c.severity)
        p.add_run(c.concern).bold = True
        _detail(doc, f"Why they may care: {c.why_they_may_care}")
        _kv(doc, "Mitigation", c.mitigation_strategy)
        _detail(doc, f"Best proof: {c.best_proof_to_use}")

    # 8. Resume Tailoring
    _section(doc, "8  Resume Tailoring")
    rt = report.resume_tailoring
    _kv(doc, "Verdict", rt.resume_verdict.replace("_", " ").title())
    if rt.positioning_angle:
        _kv(doc, "Positioning", rt.positioning_angle)
    if rt.priority_edits:
        _sub(doc, "Priority Edits")
        for edit in rt.priority_edits:
            p = doc.add_paragraph()
            p.add_run(f"[{edit.type.replace('_', ' ').title()}] {edit.target}").bold = True
            _detail(doc, edit.recommendation)
    if rt.missing_keywords:
        _kv(doc, "Missing Keywords", ", ".join(rt.missing_keywords))

    # 9. Application Strategy
    _section(doc, "9  Application Strategy")
    app_s = report.application_strategy
    _kv(doc, "Apply Now", "Yes" if app_s.apply_now else "No")
    _kv(doc, "Referral", "Recommended" if app_s.referral_recommended else "Not required")
    if app_s.outreach_angle:
        _kv(doc, "Outreach Angle", app_s.outreach_angle)
    if app_s.suggested_connection_note:
        p = doc.add_paragraph()
        run = p.add_run(f'"{app_s.suggested_connection_note}"')
        run.italic = True
        run.font.color.rgb = MUTED

    # 10. Recruiter Screen Prep
    _section(doc, "10  Recruiter Screen Prep")
    rp = report.recruiter_screen_prep
    _sub(doc, "Screening Focus Areas")
    for s in rp.what_they_will_likely_screen_for:
        doc.add_paragraph(s, style="List Bullet")
    _sub(doc, "Likely Questions")
    for q in rp.likely_recruiter_questions:
        _bold_para(doc, q.question)
        _detail(doc, f"Intent: {q.intent}")
        for pt in q.suggested_answer_points:
            doc.add_paragraph(pt, style="List Bullet")

    # 11. Story Bank
    _section(doc, "11  Story Bank")
    for s in report.story_bank:
        _bold_para(doc, s.story_title)
        _kv(doc, "Experience", s.recommended_experience)
        _kv(doc, "Angle", s.angle_to_emphasize)
        if s.key_metrics_or_outcomes:
            _kv(doc, "Metrics", ", ".join(s.key_metrics_or_outcomes))

    # 12. Interview Rounds
    _section(doc, "12  Interview Round Strategy")
    for rd in report.interview_rounds:
        _bold_para(doc, rd.round_name.replace("_", " ").title())
        doc.add_paragraph(rd.round_goal)
        for f in rd.candidate_focus:
            doc.add_paragraph(f, style="List Bullet")

    # 13. Interview Questions
    _section(doc, "13  Likely Interview Questions")
    for q in report.likely_interview_questions:
        p = doc.add_paragraph()
        cat_run = p.add_run(f"[{q.category.replace('_', ' ').title()}] ")
        cat_run.bold = True
        cat_run.font.color.rgb = INDIGO
        p.add_run(q.question).bold = True
        _detail(doc, q.why_they_may_ask)
        if q.best_story_or_topic:
            _detail(doc, f"Best story/topic: {q.best_story_or_topic}")

    # 14. Reverse Interview Questions
    _section(doc, "14  Reverse-Interview Questions")
    for q in report.reverse_interview_questions:
        p = doc.add_paragraph()
        cat_run = p.add_run(f"[{q.target_interviewer.replace('_', ' ').title()}] ")
        cat_run.bold = True
        cat_run.font.color.rgb = INDIGO
        p.add_run(q.question).bold = True
        _detail(doc, q.why_ask_this)

    # 15. Logistics
    _section(doc, "15  Logistics & Constraints")
    lg = report.logistics_and_constraints
    _kv(doc, "Work Model", lg.work_model.title())
    _kv(doc, "Level Clarity", lg.level_clarity.title())
    if lg.location_notes:
        _kv(doc, "Location", lg.location_notes)
    if lg.compensation_signal:
        _kv(doc, "Compensation", lg.compensation_signal)
    if lg.timeline_signal:
        _kv(doc, "Timeline", lg.timeline_signal)

    # 16. Red Flags
    _section(doc, "16  Red Flags & Unknowns")
    for f in report.red_flags_and_unknowns.red_flags:
        p = doc.add_paragraph()
        sev_run = p.add_run(f"[{f.severity.upper()}] ")
        sev_run.bold = True
        sev_run.font.color.rgb = _severity_color(f.severity)
        p.add_run(f.flag).bold = True
        _detail(doc, f.why_it_matters)
    if report.red_flags_and_unknowns.unknowns_to_verify:
        _sub(doc, "Unknowns to Verify")
        for u in report.red_flags_and_unknowns.unknowns_to_verify:
            doc.add_paragraph(u, style="List Bullet")

    # 17. Next Actions
    _section(doc, "17  Immediate Next Actions")
    for a in sorted(report.immediate_next_actions, key=lambda x: x.priority):
        time_str = f"  (~{a.time_estimate_minutes} min)" if a.time_estimate_minutes else ""
        p = doc.add_paragraph()
        p.add_run(f"{a.priority}.  {a.action}{time_str}").bold = True
        _detail(doc, a.why)

    # 18. Checklist
    _section(doc, "18  Preparation Checklist")
    for item in report.prep_checklist:
        doc.add_paragraph(f"☐  {item.item}")

    # ── SCORING APPENDIX ──
    if report.scoring:
        _build_scoring_appendix(doc, report.scoring)

    # ── FOOTER ──
    doc.add_paragraph()
    _add_divider(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by PrepLens  ·  Confidential  ·  For candidate use only")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Helper functions ──

def _section(doc: Document, text: str):
    """Add a numbered section heading with indigo color and bottom border feel."""
    doc.add_paragraph()  # spacing
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = INDIGO
        run.font.size = Pt(13)


def _sub(doc: Document, text: str):
    """Add a sub-heading."""
    heading = doc.add_heading(text, level=3)
    for run in heading.runs:
        run.font.color.rgb = DARK
        run.font.size = Pt(10.5)


def _bold_para(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK


def _kv(doc: Document, key: str, value: str):
    p = doc.add_paragraph()
    k_run = p.add_run(f"{key}: ")
    k_run.bold = True
    k_run.font.color.rgb = DARK
    v_run = p.add_run(value)
    v_run.font.color.rgb = DARK


def _detail(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def _add_spacing(paragraph, before: int = 0, after: int = 4):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_divider(doc: Document):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = LIGHT_BORDER


# ── Category label mapping ──
_CAT_LABELS_SHORT = {
    "hard_skill": "Hard Skill",
    "soft_skill": "Soft Skill",
    "domain_knowledge": "Domain",
    "tool_or_technology": "Tool/Tech",
    "education": "Education",
}

_MATCH_COLORS = {"direct": EMERALD, "partial": AMBER, "none": RED}
_MATCH_ICONS = {"direct": "DIRECT", "partial": "PARTIAL", "none": "NONE"}


def _build_scoring_appendix(doc: Document, scoring: dict):
    """Build the Score Methodology appendix in the DOCX."""

    doc.add_page_break()
    _add_divider(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Appendix: Score Methodology")
    run.font.size = Pt(22)
    run.font.color.rgb = INDIGO
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Transparent, ATS-style scoring breakdown — every score is traceable to evidence."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    _add_divider(doc)

    # -- Overall Score --
    _section(doc, "Overall Score")
    fit_score = scoring.get("overall_fit_score", 0)
    fit_pct = scoring.get("overall_fit_percentage", 0)
    earned = scoring.get("total_earned_points", 0)
    max_pts = scoring.get("total_max_points", 0)

    p = doc.add_paragraph()
    run = p.add_run(f"Overall Fit: {fit_score}/10")
    run.bold = True
    run.font.color.rgb = _score_color(fit_score)
    p.add_run(f"  ({fit_pct}% match)").font.color.rgb = DARK

    _kv(doc, "Total Points", f"{earned} earned out of {max_pts} possible")
    _detail(doc,
        "Score = weighted sum of category scores. Each category is scored independently "
        "based on keyword matching between the job description and resume, then weighted "
        "by the category's importance for this specific role."
    )

    # -- Confidence --
    confidence = scoring.get("confidence", {})
    if confidence:
        _section(doc, "Confidence Score")
        conf_overall = confidence.get("overall_confidence", 0)
        p = doc.add_paragraph()
        run = p.add_run(f"Overall Confidence: {conf_overall}/10")
        run.bold = True
        run.font.color.rgb = _score_color(conf_overall)

        for factor, label in [
            ("jd_specificity", "Job Description Specificity"),
            ("resume_detail", "Resume Detail Level"),
            ("match_clarity", "Match Clarity"),
        ]:
            score_val = confidence.get(factor, 0)
            reasoning = confidence.get(f"{factor}_reasoning", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: {score_val}/10")
            run.bold = True
            run.font.color.rgb = _score_color(score_val)
            _detail(doc, reasoning)

    # -- Category Breakdown Table --
    category_scores = scoring.get("category_scores", [])
    if category_scores:
        _section(doc, "Category Breakdown")
        _detail(doc,
            "Each category is weighted based on what the job description emphasizes. "
            "The weighted contribution shows how much each category affects the overall score."
        )

        table = doc.add_table(rows=1 + len(category_scores), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["Category", "Score", "Match %", "Weight", "Contribution"]
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Light background for header
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): "E0E7FF",
                qn("w:val"): "clear",
            })
            shading.append(shading_elm)

        for row_idx, cs in enumerate(category_scores, start=1):
            vals = [
                cs.get("category_label", cs["category"]),
                f'{cs["score_out_of_10"]}/10',
                f'{cs["percentage"]}%',
                f'{cs["category_weight"]}',
                f'{cs["weighted_contribution"]}',
            ]
            for col_idx, val in enumerate(vals):
                cell = table.cell(row_idx, col_idx)
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                if col_idx == 1:
                    run.font.color.rgb = _score_color(cs["score_out_of_10"])
                    run.bold = True
                else:
                    run.font.color.rgb = DARK
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # -- Requirements Breakdown Table --
    scored_reqs = scoring.get("scored_requirements", [])
    if scored_reqs:
        _section(doc, "Detailed Requirements Breakdown")
        _detail(doc,
            "Every requirement extracted from the job description, matched against your resume. "
            "Weight: Required (5), Preferred (3), Nice-to-have (1)."
        )

        table = doc.add_table(rows=1 + len(scored_reqs), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Requirement", "Category", "Importance", "Match", "Points"]
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): "E0E7FF",
                qn("w:val"): "clear",
            })
            shading.append(shading_elm)

        for row_idx, req in enumerate(scored_reqs, start=1):
            m_level = req.get("match_level", "none")
            vals = [
                req["keyword"],
                _CAT_LABELS_SHORT.get(req["category"], req["category"]),
                req["importance"].replace("_", " ").title(),
                _MATCH_ICONS.get(m_level, m_level),
                f'{req["earned_points"]}/{req["max_points"]}',
            ]
            for col_idx, val in enumerate(vals):
                cell = table.cell(row_idx, col_idx)
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                if col_idx == 3:
                    run.font.color.rgb = _MATCH_COLORS.get(m_level, MUTED)
                    run.bold = True
                else:
                    run.font.color.rgb = DARK
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # -- Match Evidence --
    if scored_reqs:
        _section(doc, "Match Evidence")
        _detail(doc, "Detailed justification for each requirement match.")

        for req in scored_reqs:
            m_level = req.get("match_level", "none")
            p = doc.add_paragraph()
            run = p.add_run(f'{req["keyword"]} — ')
            run.bold = True
            run.font.color.rgb = DARK
            run2 = p.add_run(m_level.upper())
            run2.bold = True
            run2.font.color.rgb = _MATCH_COLORS.get(m_level, MUTED)
            _detail(doc, f"Evidence: {req.get('resume_evidence', 'N/A')}")
            _detail(doc, f"Reasoning: {req.get('match_reasoning', 'N/A')}")

    # -- Top Improvements --
    improvements = scoring.get("top_improvements", [])
    if improvements:
        _section(doc, "Top Score Improvements")
        _detail(doc, "Changes that would have the biggest impact on your overall score.")

        for imp in improvements:
            p = doc.add_paragraph()
            run = p.add_run(
                f'{imp["keyword"]} — +{imp["potential_points"]} pts '
                f'({imp["percentage_impact"]}% impact)'
            )
            run.bold = True
            run.font.color.rgb = DARK
            _detail(doc, imp["suggestion"])
